import streamlit as st
from openai import OpenAI
import os
from docx import Document
from docx.shared import Pt
import io
import re
from dotenv import load_dotenv
load_dotenv()

client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])

def model_list_conversion(user_input):
    response = client.chat.completions.create(
    model="gpt-3.5-turbo",
    messages=[
        {
            "role": "system",
            "content" : "You are a helpful assistant. For once, please provide the input you want to be converted into a list and also Omiting the un wanted comma"
        },
        {
        "role": "user",
        "content": "The details for the transaction with invoice number 52015342 are as follows:\n\nCustomer Name: Jeffrey Johnson\nEmail: mchen@example.net\nTransaction Amount: 39.04 USD\nAddress: 783 Owens Way\nCity: Lake Ashleyborough\nInvoice Number: 52015342.0\nTransaction Date: 25-07-2021\nProduct Name: 7 For All Mankind Men's Standard Classic Straight Leg Jean In Authentic Nakita\nCompany Code: DSAI895623\nShould you need assistance with anything else, feel free to ask."
        },
        {
        "role": "assistant",
        "content": "Jeffrey Johnson, mchen@example.net, 39.04 USD, 783 Owens Way, Lake Ashleyborough, 52015342, 25-07-2021, 7 For All Mankind Men's Standard Classic Straight Leg Jean, DSAI895623"
        },
        {
            "role":"user",
            "content": user_input
        }
    ],
    temperature=0,
    max_tokens=500,
    top_p=1,
    frequency_penalty=0,
    presence_penalty=0
    )
    assistant_response = response.choices[0].message.content
    return assistant_response

def get_invoice_details(thread_id,user_input):
    user_input = f'My invoice number is {user_input}'
    client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
    message_1 = client.beta.threads.messages.create(thread_id=thread_id,role="user",content=user_input)
    run = client.beta.threads.runs.create(thread_id=thread_id,assistant_id=os.environ["CALL_CENTER_CALLBOT_TRANSACTION_REGISTER"])
    # run = client.beta.threads.runs.create(thread_id=thread.id,assistant_id=assistant.id,run_id=run.id)
    while True:
        run = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run.id)
        if run.status=="completed":
            messages = client.beta.threads.messages.list(thread_id=thread_id)
            latest_message = messages.data[0]
            text = latest_message.content[0].text.value
            break
    return model_list_conversion(str(text))

def writting_in_doc(values_list):
    # Load your existing Word document
    doc = Document('invoice/Invoice.docx')
    # Assuming the table is the first table in the document
    table = doc.tables[0]

    # Set the font style for the second column
    font_style = 'Century Gothic'

    # Iterate through rows and update the second column
    for i, value in enumerate(values_list):
        cell = table.cell(i, 1)
        cell.text = value  # Column index is 1 for the second column (0-based index)

        # Set font style
        run = cell.paragraphs[0].runs[0]
        font = run.font
        font.name = font_style
        font.size = Pt(12)  # You can adjust the font size as needed
    file_path = 'invoice/user_Invoice.docx'
    
    # Save the modified document
    doc.save(file_path)

def Transaction_invoice_download(thread_id):
    col1,col2,col3,col4 = st.columns((2,2.5,3.5,2))
    with col2:
        st.write('# ')
        st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: bold'>Invoice Number</span></p>", unsafe_allow_html=True)
    with col3:
        vAR_invoice_num = st.text_input(" ")
    if vAR_invoice_num:
        vAR_invoice_as_list = get_invoice_details(thread_id,vAR_invoice_num)
        
        result_list = [item.strip() for item in vAR_invoice_as_list.split(',')]
        
        if len(result_list) == 9:
            writting_in_doc(result_list)
            file_path = 'invoice/user_Invoice.docx'
            doc = Document(file_path)
            with col3:
                st.markdown("")
                bio = io.BytesIO()
                doc.save(bio)
                vAR_invoice_num = ""
                st.download_button(
                    label="Download",
                    data = bio.getvalue(),
                    file_name= 'user_Invoice.docx',
                    mime="docx"
                )
        else:
            with col3:
                st.warning("Try Again")